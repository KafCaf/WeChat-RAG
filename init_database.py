import os
import fitz  # 处理 PDF 的 PyMuPDF 库
import pandas as pd  # 处理 Excel 的 Pandas 库
from datetime import datetime
from docx import Document
from retrievers.VectorRetriever import VectorRetrieval
from configs.model_configs import EMBED_CONFIG
from utils import list_files_from_folder, list_kbs_from_folder
from pathlib import Path
from elasticsearch import Elasticsearch, helpers
from server.embedding import CloudEmbedModel
from configs.model_configs import MODEL_PATH, SPLITTER_CONFIG
from text_splitter.TSdocx_splitter import TSDocTextSplitter
import warnings
import traceback
import re

warnings.filterwarnings("ignore")
retrieval_classes = {
    "vector": VectorRetrieval
}

# ==================== 1. 统一文本提取层 (路由分发) ====================
def extract_text_from_file(file_path):
    """根据文件后缀，调用对应的解析器提取文本段落，统一返回 full_text_list"""
    ext = os.path.splitext(file_path)[1].lower()
    full_text_list = []
    
    print(f"正在读取并解析文件: {file_path} (格式: {ext})")

    try:
        # ---- [格式 1]: DOCX 文本与表格解析 ----
        if ext == '.docx':
            doc = Document(file_path)

            try:
                # 1. 找到并彻底销毁所有带有 <w:del> 标签的节点（彻底删除不要的废弃文本）
                for del_node in doc.element.xpath('//w:del'):
                    del_node.getparent().remove(del_node)
                
                # 2. 找到所有带有 <w:ins> 标签的新增节点，把里面的文字“释放”出来
                for ins_node in doc.element.xpath('//w:ins'):
                    parent = ins_node.getparent()
                    index = parent.index(ins_node)
                    # 将隐藏在 <w:ins> 内部的有效文本节点剥离出来，挂载到正常的段落树上
                    for child in list(ins_node):
                        parent.insert(index, child)
                        index += 1
                    # 删除没用的 <w:ins> 外壳
                    parent.remove(ins_node)
            except Exception as e:
                print(f"清洗修订标记时出现小问题，但不影响继续解析: {e}")
            
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text_list.append(para.text.strip())
            
            # 提取表格内容（通用列标注，不假设表格语义）
            if len(doc.tables) > 0:
                for table in doc.tables:
                    if not table.rows:
                        continue
                    # 提取表头
                    header_cells = [cell.text.strip().replace("\n", "") for cell in table.rows[0].cells]
                    header_count = len(header_cells)
                    for i, row in enumerate(table.rows):
                        if i == 0:
                            continue
                        row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        if not any(row_cells):
                            continue
                        # 补齐合并单元格导致的列数不足
                        while len(row_cells) < header_count:
                            row_cells.append("")
                        parts = []
                        for idx, cell_text in enumerate(row_cells):
                            if cell_text:
                                h = header_cells[idx] if idx < header_count else f"列{idx+1}"
                                parts.append(f"【{h}】{cell_text}")
                        full_text_list.append("  ".join(parts))

        # ---- [格式 2]: PDF 解析 ----
        elif ext == '.pdf':
            with fitz.open(file_path) as doc:
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text = page.get_text("text")
                    if text.strip():
                        full_text_list.append(text.strip())

        # ---- [格式 3 & 4]: 纯文本类 (TXT / Markdown) ----
        elif ext in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                if content.strip():
                    # 将纯文本按一定的换行符稍微分割，避免单行过长
                    paragraphs = content.split('\n\n')
                    for p in paragraphs:
                        if p.strip(): full_text_list.append(p.strip())

        # ---- [格式 5]: EXCEL 表格降维解析 ----
        elif ext in ['.xlsx', '.xls']:
            # 将 Excel 的每一行转化为结构化的语义文本
            excel_data = pd.read_excel(file_path, sheet_name=None)
            for sheet_name, df in excel_data.items():
                full_text_list.append(f"--- 表格所在工作表：{sheet_name} ---")
                df = df.fillna("") # 将空值转换为占位符
                headers = list(df.columns)
                for index, row in df.iterrows():
                    row_content_parts = []
                    for col_name in headers:
                        cell_value = str(row[col_name]).strip()
                        if cell_value:
                            row_content_parts.append(f"【{col_name}】: {cell_value}")
                    if row_content_parts:
                        full_text_list.append("  ".join(row_content_parts))

        else:
            print(f"系统提示：当前暂不支持解析 {ext} 格式文件，已跳过。")

    except Exception as e:
        print(f"提取文件 {file_path} 内容时崩溃: {e}")
        traceback.print_exc()

    return full_text_list


# ==================== 2. 统一分块层 ====================
def process_document(file_path):
    """解析文档并分块。.docx 文件使用标题感知切分，其他格式使用固定窗口。"""
    chunks = {}
    ext = os.path.splitext(file_path)[1].lower()
    project_name = Path(file_path).parent.name
    
    # 步骤一：提取纯文本列表
    full_text_list = extract_text_from_file(file_path)
    if not full_text_list:
        return None
    
    # 步骤二：通用标题感知切分
    chunk_list = []
    splitter = TSDocTextSplitter(chunk_size=1200)
    
    if ext == '.docx':
        # 先尝试 docx 样式+文本双重检测
        try:
            doc = Document(file_path)
            heading_chunks = splitter.split_text(doc)
            for c in heading_chunks:
                if len(c) > 50:
                    chunk_list.append(f"【来源政策项目：{project_name}】\n" + c)
            if chunk_list:
                print(f"文件 {Path(file_path).name} 标题感知切片完成，共生成 {len(chunk_list)} 个数据块。")
        except Exception as e:
            print(f"标题感知切分失败 ({e})，回退通用模式")
    
    if not chunk_list:
        # 回退：对全部提取文本（段落+表格行）做通用标题检测
        chunk_list = splitter.split_lines(full_text_list)
        chunk_list = [f"【来源政策项目：{project_name}】\n" + c for c in chunk_list if len(c) > 50]
        print(f"文件 {Path(file_path).name} 通用模式切片完成，共生成 {len(chunk_list)} 个数据块。")

    chunks["text"] = chunk_list 
    chunks["filename"] = str(Path(file_path).as_posix())
    chunks["project_name"] = project_name
    chunks["date"] = datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%S')
    
    return chunks


# ==================== 3. 数据库交互与遍历层 ====================
def files2db(retrieval_type, es_client, index_name, kb_files, embed_model):
    kb_service = retrieval_classes[retrieval_type](embed_model=embed_model, es_client=es_client)
    for kb_file in kb_files:
        try:
            # 统一入口改为了 process_document
            chunks = process_document(kb_file) 
            if chunks and chunks.get("text"):
                print(f"正在将 {Path(kb_file).name} 注入知识库 (ES Index: {index_name})...")
                kb_service.build_index(index_name, chunks)
            else:
                print(f"警告：文件 {kb_file} 解析内容为空，已跳过注入。")
                continue
        except Exception as e:
            print(f"处理文件报错 {kb_file}: {e}", flush=True)
            traceback.print_exc()
            continue


def folder2db(kb_names, retrieval_type, es_client, index_name, embed_model):
    kb_names = kb_names or list_kbs_from_folder()
    
    for kb_name in kb_names:
        kb_files = list_files_from_folder(kb_name)
        files2db(retrieval_type, es_client, index_name, kb_files, embed_model)

# ----------------- ES 与模型初始化 -----------------
if __name__ == "__main__":
    es_client = Elasticsearch("http://es:9200")
    print(es_client.info())

    if es_client.ping():
        print("Successfully connected to Elasticsearch!")
        embed_model = CloudEmbedModel()
    else:
        print("Could not connect to Elasticsearch.")
        exit(1)

    index_name = "index_user_test"
    try:
        es_client.indices.delete(index=index_name, ignore=[400, 404])
        print(f"旧索引 {index_name} 已清理/删除，准备重新构建...")
    except Exception as e:
        print(f"删除索引时遇到问题 (可能是连接问题): {e}")
        
    folder2db(None, retrieval_type="vector", es_client=es_client, index_name="index_user_test", embed_model=embed_model)
